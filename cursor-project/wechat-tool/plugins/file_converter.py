"""
文件转换插件
实现Word、PDF、图片之间的转换功能
"""

import os
import shutil
from pathlib import Path
from typing import Optional, Tuple
from loguru import logger

import itchat
from docx import Document
from docx2txt import process
import PyPDF2
from PIL import Image
import img2pdf
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter, A4

from core.base_plugin import BasePlugin


class FileConverterPlugin(BasePlugin):
    """文件转换插件"""
    
    def __init__(self, config_manager):
        super().__init__(config_manager)
        self.name = "FileConverter"
        self.version = "1.0.0"
        self.description = "文件格式转换工具，支持Word、PDF、图片互转"
        
        # 获取配置
        self.upload_dir = self.config.get("file_processing.upload_dir", "uploads")
        self.output_dir = self.config.get("file_processing.output_dir", "outputs")
        self.max_file_size = self.config.get("file_processing.max_file_size", 50) * 1024 * 1024  # MB to bytes
        self.supported_formats = self.config.get("file_processing.supported_formats", {})
        
        # 确保目录存在
        Path(self.upload_dir).mkdir(exist_ok=True)
        Path(self.output_dir).mkdir(exist_ok=True)
    
    def start(self) -> None:
        """启动插件"""
        self.is_running = True
        logger.info("文件转换插件已启动")
    
    def stop(self) -> None:
        """停止插件"""
        self.is_running = False
        logger.info("文件转换插件已停止")
    
    def handle_file_upload(self, file_path: str, file_name: str, user_id: str) -> str:
        """处理文件上传"""
        try:
            logger.info(f"处理文件上传: {file_name}")
            
            # 检查文件大小
            if not self._check_file_size(file_path):
                return f"文件过大，最大支持 {self.max_file_size // (1024*1024)}MB"
            
            # 获取文件类型
            file_type = self._get_file_type(file_name)
            if not file_type:
                return "不支持的文件格式"
            
            # 根据文件类型进行转换
            if file_type == "word":
                return self._convert_word_to_pdf(file_path, file_name, user_id)
            elif file_type == "image":
                return self._convert_image_to_pdf(file_path, file_name, user_id)
            elif file_type == "pdf":
                return self._convert_pdf_to_word(file_path, file_name, user_id)
            else:
                return "不支持的文件格式"
                
        except Exception as e:
            logger.error(f"处理文件上传失败: {e}")
            return f"文件处理失败: {str(e)}"
    
    def execute_command(self, command: str, **kwargs) -> str:
        """执行插件命令"""
        if command == "help":
            return self.get_help()
        elif command == "status":
            return self._get_conversion_status()
        else:
            return "未知命令，输入 'help' 查看帮助"
    
    def get_help(self) -> str:
        """获取帮助信息"""
        return """
📁 文件转换插件帮助

🔄 支持的转换：
• Word文档 (.doc/.docx) → PDF
• 图片文件 (.jpg/.png/.bmp等) → PDF  
• PDF文件 → Word文档

📋 使用方法：
• 直接发送文件即可自动转换
• 命令: /file_converter help - 显示帮助
• 命令: /file_converter status - 查看状态

💡 提示：
• 支持的文件大小: 最大50MB
• 转换后的文件会自动发送给您
        """
    
    def _check_file_size(self, file_path: str) -> bool:
        """检查文件大小"""
        try:
            file_size = os.path.getsize(file_path)
            return file_size <= self.max_file_size
        except Exception:
            return False
    
    def _get_file_type(self, file_name: str) -> Optional[str]:
        """获取文件类型"""
        file_ext = Path(file_name).suffix.lower()
        
        if file_ext in self.supported_formats.get("word", []):
            return "word"
        elif file_ext in self.supported_formats.get("pdf", []):
            return "pdf"
        elif file_ext in self.supported_formats.get("image", []):
            return "image"
        else:
            return None
    
    def _convert_word_to_pdf(self, file_path: str, file_name: str, user_id: str) -> str:
        """Word转PDF"""
        try:
            logger.info(f"开始转换Word到PDF: {file_name}")
            
            # 读取Word文档
            doc = Document(file_path)
            
            # 创建输出文件名
            output_name = Path(file_name).stem + ".pdf"
            output_path = os.path.join(self.output_dir, output_name)
            
            # 提取文本内容
            text_content = process(file_path)
            
            # 创建PDF
            c = canvas.Canvas(output_path, pagesize=A4)
            width, height = A4
            
            # 设置字体和大小
            c.setFont("Helvetica", 12)
            
            # 分行处理文本
            lines = text_content.split('\n')
            y_position = height - 50
            
            for line in lines:
                if y_position < 50:  # 需要新页面
                    c.showPage()
                    c.setFont("Helvetica", 12)
                    y_position = height - 50
                
                # 处理长行
                if len(line) > 80:
                    words = line.split()
                    current_line = ""
                    for word in words:
                        if len(current_line + word) < 80:
                            current_line += word + " "
                        else:
                            c.drawString(50, y_position, current_line.strip())
                            y_position -= 20
                            current_line = word + " "
                    
                    if current_line:
                        c.drawString(50, y_position, current_line.strip())
                        y_position -= 20
                else:
                    c.drawString(50, y_position, line)
                    y_position -= 20
            
            c.save()
            
            # 发送转换后的文件
            itchat.send_file(output_path, toUserName=user_id)
            
            # 清理临时文件
            os.remove(output_path)
            
            return f"✅ Word文档已成功转换为PDF！"
            
        except Exception as e:
            logger.error(f"Word转PDF失败: {e}")
            return f"转换失败: {str(e)}"
    
    def _convert_image_to_pdf(self, file_path: str, file_name: str, user_id: str) -> str:
        """图片转PDF"""
        try:
            logger.info(f"开始转换图片到PDF: {file_name}")
            
            # 创建输出文件名
            output_name = Path(file_name).stem + ".pdf"
            output_path = os.path.join(self.output_dir, output_name)
            
            # 使用img2pdf转换
            with open(output_path, "wb") as f:
                f.write(img2pdf.convert(file_path))
            
            # 发送转换后的文件
            itchat.send_file(output_path, toUserName=user_id)
            
            # 清理临时文件
            os.remove(output_path)
            
            return f"✅ 图片已成功转换为PDF！"
            
        except Exception as e:
            logger.error(f"图片转PDF失败: {e}")
            return f"转换失败: {str(e)}"
    
    def _convert_pdf_to_word(self, file_path: str, file_name: str, user_id: str) -> str:
        """PDF转Word"""
        try:
            logger.info(f"开始转换PDF到Word: {file_name}")
            
            # 读取PDF
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                # 提取文本
                text_content = ""
                for page in pdf_reader.pages:
                    text_content += page.extract_text() + "\n"
            
            # 创建Word文档
            doc = Document()
            
            # 添加文本内容
            paragraphs = text_content.split('\n')
            for para in paragraphs:
                if para.strip():
                    doc.add_paragraph(para.strip())
            
            # 创建输出文件名
            output_name = Path(file_name).stem + ".docx"
            output_path = os.path.join(self.output_dir, output_name)
            
            # 保存Word文档
            doc.save(output_path)
            
            # 发送转换后的文件
            itchat.send_file(output_path, toUserName=user_id)
            
            # 清理临时文件
            os.remove(output_path)
            
            return f"✅ PDF已成功转换为Word文档！"
            
        except Exception as e:
            logger.error(f"PDF转Word失败: {e}")
            return f"转换失败: {str(e)}"
    
    def _get_conversion_status(self) -> str:
        """获取转换状态"""
        return f"""
📊 文件转换插件状态

✅ 运行状态: {'运行中' if self.is_running else '已停止'}
📁 上传目录: {self.upload_dir}
📁 输出目录: {self.output_dir}
📏 最大文件大小: {self.max_file_size // (1024*1024)}MB

🔄 支持的转换:
• Word → PDF: ✅
• 图片 → PDF: ✅  
• PDF → Word: ✅
        """ 